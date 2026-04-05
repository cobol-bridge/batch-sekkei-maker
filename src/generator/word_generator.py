"""
Word形式の設計書生成
"""
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# テーマカラー（Excelと統一）
HEADER_BG = "1F4E79"   # 紺
HEADER_FG = "FFFFFF"   # 白
SUBHEADER_BG = "BDD7EE"  # 薄青


def _set_cell_bg(cell, hex_color: str):
    """セルの背景色を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell):
    """セルに細枠線を設定"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _header_cell(cell, text: str):
    """ヘッダーセルのスタイルを適用"""
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_bg(cell, HEADER_BG)
    _set_cell_border(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0] if para.runs else para.add_run(text)
    run.text = text
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "メイリオ"
    run.font.size = Pt(9)


def _body_cell(cell, text: str, bold: bool = False, bg: str = None):
    """ボディセルのスタイルを適用"""
    cell.text = str(text) if text else ""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if bg:
        _set_cell_bg(cell, bg)
    _set_cell_border(cell)
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run(str(text) if text else "")
    run.text = str(text) if text else ""
    run.font.name = "メイリオ"
    run.font.size = Pt(9)
    run.font.bold = bold


def _section_heading(doc, text: str):
    """セクション見出しを追加"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "メイリオ"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def generate_word(parse_result, output_path: str):
    """
    パース結果からWord設計書を生成する

    Args:
        parse_result: CobolParseResult
        output_path: 出力先ファイルパス（.docx）
    """
    doc = Document()

    # 用紙をA4横向きに設定
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    _add_cover(doc, parse_result)
    _add_io_section(doc, parse_result)
    _add_flow_section(doc, parse_result)
    _add_exception_section(doc, parse_result)

    doc.save(output_path)


def _add_cover(doc, parse_result):
    """表紙セクション"""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(20)
    title.paragraph_format.space_after = Pt(20)
    run = title.add_run("バッチ処理 業務設計書")
    run.font.name = "メイリオ"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    rows_data = [
        ("対象ファイル", parse_result.source_file),
        ("生成日時", datetime.now().strftime("%Y年%m月%d日 %H:%M")),
        ("ツールバージョン", "v1.0.0"),
        ("備考", "本設計書はバッチ設計書メーカーにより自動生成されました。\n内容は必ず目視で確認してください。"),
    ]

    col_widths = [Cm(4), Cm(20)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    for i, (label, value) in enumerate(rows_data):
        _body_cell(table.rows[i].cells[0], label, bold=True, bg=SUBHEADER_BG)
        _body_cell(table.rows[i].cells[1], value)

    doc.add_paragraph()


def _add_io_section(doc, parse_result):
    """I/O定義セクション"""
    _section_heading(doc, "I/O定義一覧")

    headers = ["No", "ファイル名（SELECT）", "FD名", "レコード長",
               "物理ファイル名", "ファイル編成", "アクセスモード", "レコードキー"]
    col_widths = [Cm(1.0), Cm(3.5), Cm(3.5), Cm(2.0),
                  Cm(3.5), Cm(2.8), Cm(2.8), Cm(3.5)]

    if not parse_result.file_definitions:
        doc.add_paragraph("（I/O定義が検出されませんでした）")
        return

    table = doc.add_table(rows=1 + len(parse_result.file_definitions), cols=len(headers))

    for i, (w, h) in enumerate(zip(col_widths, headers)):
        cell = table.rows[0].cells[i]
        cell.width = w
        _header_cell(cell, h)

    for row_idx, fd in enumerate(parse_result.file_definitions, start=1):
        cells = table.rows[row_idx].cells
        values = [
            str(row_idx),
            fd.file_name or "（不明）",
            fd.fd_name or "（不明）",
            fd.record_length or "（不明）",
            fd.assign_to or "",
            fd.organization or "",
            fd.access_mode or "",
            fd.record_key or "",
        ]
        for i, (val, w) in enumerate(zip(values, col_widths)):
            cells[i].width = w
            _body_cell(cells[i], val)


def _add_flow_section(doc, parse_result):
    """処理フローセクション"""
    _section_heading(doc, "処理フロー")

    headers = ["No", "呼び出し元（段落名）", "呼び出し先（段落名）", "種別", "備考"]
    col_widths = [Cm(1.0), Cm(5.5), Cm(5.5), Cm(3.5), Cm(6.0)]

    if not parse_result.perform_entries:
        doc.add_paragraph("（処理フローが検出されませんでした）")
        return

    table = doc.add_table(rows=1 + len(parse_result.perform_entries), cols=len(headers))

    for i, (w, h) in enumerate(zip(col_widths, headers)):
        cell = table.rows[0].cells[i]
        cell.width = w
        _header_cell(cell, h)

    for row_idx, entry in enumerate(parse_result.perform_entries, start=1):
        cells = table.rows[row_idx].cells
        values = [str(row_idx), entry.caller or "（不明）", entry.callee, entry.perform_type, ""]
        for i, (val, w) in enumerate(zip(values, col_widths)):
            cells[i].width = w
            _body_cell(cells[i], val)


def _add_exception_section(doc, parse_result):
    """例外処理セクション"""
    _section_heading(doc, "例外処理一覧")

    headers = ["No", "発生箇所（段落名）", "例外種別", "処理内容（抜粋）", "備考"]
    col_widths = [Cm(1.0), Cm(5.0), Cm(4.0), Cm(8.0), Cm(3.5)]

    if not parse_result.exception_entries:
        doc.add_paragraph("（例外処理が検出されませんでした）")
        return

    table = doc.add_table(rows=1 + len(parse_result.exception_entries), cols=len(headers))

    for i, (w, h) in enumerate(zip(col_widths, headers)):
        cell = table.rows[0].cells[i]
        cell.width = w
        _header_cell(cell, h)

    for row_idx, entry in enumerate(parse_result.exception_entries, start=1):
        cells = table.rows[row_idx].cells
        values = [str(row_idx), entry.location or "（不明）", entry.exception_type, entry.handling, ""]
        for i, (val, w) in enumerate(zip(values, col_widths)):
            cells[i].width = w
            _body_cell(cells[i], val)
