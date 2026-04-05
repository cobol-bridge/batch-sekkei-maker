"""
COBOLパーサーのユニットテスト
"""
import pytest
import os
from src.parser.cobol_parser import CobolParser


SAMPLE_FILE = os.path.join(os.path.dirname(__file__), "../samples/SAMPLE01.cbl")


class TestCobolParser:

    def setup_method(self):
        self.parser = CobolParser()

    def test_parse_file_exists(self):
        """サンプルファイルが存在すること"""
        assert os.path.isfile(SAMPLE_FILE), f"サンプルファイルが見つかりません: {SAMPLE_FILE}"

    def test_parse_returns_result(self):
        """パースが結果を返すこと"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        assert result is not None

    def test_parse_no_errors(self):
        """サンプルファイルでエラーが出ないこと"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        assert result.errors == [], f"エラーが発生しました: {result.errors}"

    def test_parse_file_definitions(self):
        """SELECT句からファイル定義が3件抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        assert len(result.file_definitions) == 3, \
            f"ファイル定義数が不正です（期待：3件、実際：{len(result.file_definitions)}件）"

    def test_parse_file_names(self):
        """SELECT句のファイル名が正しく抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        file_names = [fd.file_name for fd in result.file_definitions]
        assert "URIAGE-FILE" in file_names
        assert "TOKUISAKI-FILE" in file_names
        assert "SHUUKEI-FILE" in file_names

    def test_parse_organization(self):
        """ORGANIZATION IS が正しく抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        fd_map = {fd.file_name: fd for fd in result.file_definitions}
        assert fd_map["URIAGE-FILE"].organization == "SEQUENTIAL"
        assert fd_map["TOKUISAKI-FILE"].organization == "INDEXED"

    def test_parse_access_mode(self):
        """ACCESS MODE IS が正しく抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        fd_map = {fd.file_name: fd for fd in result.file_definitions}
        assert fd_map["TOKUISAKI-FILE"].access_mode == "RANDOM"

    def test_parse_record_key(self):
        """RECORD KEY IS が正しく抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        fd_map = {fd.file_name: fd for fd in result.file_definitions}
        assert fd_map["TOKUISAKI-FILE"].record_key == "TK-CODE"

    def test_no_duplicate_exception_types(self):
        """同一箇所・同一種別の例外が重複しないこと"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        seen = set()
        for e in result.exception_entries:
            key = (e.location, e.exception_type)
            assert key not in seen, f"例外が重複しています: {key}"
            seen.add(key)

    def test_parse_perform_entries(self):
        """PERFORM文が1件以上抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        assert len(result.perform_entries) > 0, "PERFORM文が検出されませんでした"

    def test_parse_exception_entries(self):
        """例外処理句が1件以上抽出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        assert len(result.exception_entries) > 0, "例外処理が検出されませんでした"

    def test_exception_types(self):
        """INVALID KEY と AT END と ON SIZE ERROR が検出されること"""
        result = self.parser.parse_file(SAMPLE_FILE, encoding="utf-8")
        types = [e.exception_type for e in result.exception_entries]
        assert "INVALID KEY" in types
        assert "AT END" in types
        assert "ON SIZE ERROR" in types

    def test_file_not_found(self):
        """存在しないファイルはエラーリストに記録されること"""
        result = self.parser.parse_file("存在しないファイル.cbl")
        assert len(result.errors) > 0


class TestExcelGenerator:

    def test_generate_excel(self, tmp_path):
        """Excelファイルが正常に生成されること"""
        from src.parser.cobol_parser import CobolParser
        from src.generator.excel_generator import generate_excel

        parser = CobolParser()
        result = parser.parse_file(SAMPLE_FILE, encoding="utf-8")

        output_path = str(tmp_path / "test_output.xlsx")
        generate_excel(result, output_path)

        assert os.path.isfile(output_path), "Excelファイルが生成されませんでした"

    def test_excel_has_4_sheets(self, tmp_path):
        """Excelファイルに4シートが存在すること"""
        import openpyxl
        from src.parser.cobol_parser import CobolParser
        from src.generator.excel_generator import generate_excel

        parser = CobolParser()
        result = parser.parse_file(SAMPLE_FILE, encoding="utf-8")

        output_path = str(tmp_path / "test_output.xlsx")
        generate_excel(result, output_path)

        wb = openpyxl.load_workbook(output_path)
        assert len(wb.sheetnames) == 4, f"シート数が不正です: {wb.sheetnames}"
        assert "表紙" in wb.sheetnames
        assert "IO定義" in wb.sheetnames
        assert "処理フロー" in wb.sheetnames
        assert "例外処理" in wb.sheetnames


class TestWordGenerator:

    def test_generate_word(self, tmp_path):
        """Wordファイルが正常に生成されること"""
        from src.parser.cobol_parser import CobolParser
        from src.generator.word_generator import generate_word

        parser = CobolParser()
        result = parser.parse_file(SAMPLE_FILE, encoding="utf-8")

        output_path = str(tmp_path / "test_output.docx")
        generate_word(result, output_path)

        assert os.path.isfile(output_path), "Wordファイルが生成されませんでした"

    def test_word_contains_sections(self, tmp_path):
        """Wordファイルに4セクションの見出しが含まれること"""
        from docx import Document
        from src.parser.cobol_parser import CobolParser
        from src.generator.word_generator import generate_word

        parser = CobolParser()
        result = parser.parse_file(SAMPLE_FILE, encoding="utf-8")

        output_path = str(tmp_path / "test_output.docx")
        generate_word(result, output_path)

        doc = Document(output_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "バッチ処理 業務設計書" in text
        assert "I/O定義一覧" in text
        assert "処理フロー" in text
        assert "例外処理一覧" in text

    def test_word_io_rows(self, tmp_path):
        """I/O定義テーブルにファイル定義の行が含まれること"""
        from docx import Document
        from src.parser.cobol_parser import CobolParser
        from src.generator.word_generator import generate_word

        parser = CobolParser()
        result = parser.parse_file(SAMPLE_FILE, encoding="utf-8")

        output_path = str(tmp_path / "test_output.docx")
        generate_word(result, output_path)

        doc = Document(output_path)
        all_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "URIAGE-FILE" in all_text
        assert "TOKUISAKI-FILE" in all_text
